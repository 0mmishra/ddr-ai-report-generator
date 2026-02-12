"""DOCX report generation for validated DDR outputs."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Union

from docx import Document
from docx.shared import Pt

from backend.schemas import DDRReport


def _set_base_style(document: Document) -> None:
    """Apply readable default typography across the document."""
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def generate_docx(ddr_report: Union[DDRReport, dict], output_dir: str = "outputs") -> str:
    """Render the DDR report as a professional DOCX and return saved file path."""
    report = ddr_report if isinstance(ddr_report, DDRReport) else DDRReport.model_validate(ddr_report)

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    file_path = output_path / f"ddr_report_{timestamp}.docx"

    doc = Document()
    _set_base_style(doc)

    doc.add_heading("Detailed Diagnostic Report", level=0)

    doc.add_heading("Property Issue Summary", level=1)
    doc.add_paragraph(report.Property_Issue_Summary)

    doc.add_heading("Area Wise Observations", level=1)
    for area_name, observation in report.Area_Wise_Observations.items():
        doc.add_heading(area_name, level=2)
        doc.add_paragraph(f"Inspection Observation: {observation.inspection_observation}")
        doc.add_paragraph(f"Thermal Observation: {observation.thermal_observation}")
        doc.add_paragraph(f"Merged Finding: {observation.merged_finding}")
        doc.add_paragraph(f"Dampness Type: {observation.dampness_type}")
        doc.add_paragraph(f"Inspection Evidence Reference: {observation.inspection_evidence_ref}")
        doc.add_paragraph(f"Thermal Evidence Reference: {observation.thermal_evidence_ref}")
        doc.add_paragraph(f"Conflict Note: {observation.conflict_note}")

    doc.add_heading("Probable Root Cause", level=1)
    doc.add_paragraph(report.Probable_Root_Cause)

    doc.add_heading("Severity Assessment", level=1)
    doc.add_paragraph(f"Overall Severity: {report.Severity_Assessment.overall_severity}")
    doc.add_paragraph(f"Reasoning: {report.Severity_Assessment.reasoning}")
    doc.add_paragraph(f"Confidence Level: {report.Severity_Assessment.Confidence_Level}")
    doc.add_paragraph(f"Confidence Reasoning: {report.Severity_Assessment.Confidence_Reasoning}")

    doc.add_heading("Recommended Actions", level=1)
    doc.add_paragraph(report.Recommended_Actions)

    doc.add_heading("Risk Implications", level=1)
    doc.add_paragraph(report.Risk_Implications)

    doc.add_heading("Additional Notes", level=1)
    doc.add_paragraph(report.Additional_Notes)

    doc.add_heading("Missing/Unclear Information", level=1)
    for item in report.Missing_or_Unclear_Information:
        doc.add_paragraph(item, style="List Bullet")

    doc.save(str(file_path))
    return str(file_path)
